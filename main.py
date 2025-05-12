from flask import Flask, request, render_template, send_from_directory, abort, after_this_request
from werkzeug.utils import secure_filename
import pandas as pd
from docx import Document
import os
import re
import zipfile
import datetime
import logging

# Configure logging
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s [%(levelname)s] %(message)s',
    handlers=[
        logging.StreamHandler()
    ]
)

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = '/tmp/uploads'
app.config['OUTPUT_FOLDER'] = '/tmp/output'
app.config['MAX_CONTENT_LENGTH'] = 5 * 1024 * 1024  # 5 MB for free hosting

# Create folders if they don't exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)

ALLOWED_WORD_EXTENSIONS = {'.docx'}
ALLOWED_EXCEL_EXTENSIONS = {'.xlsx', '.xls'}

def allowed_file(filename, allowed_extensions):
    return '.' in filename and os.path.splitext(filename)[1].lower() in allowed_extensions

@app.route('/', methods=['GET'])
def index():
    logging.info("Rendering index page")
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_files():
    logging.info("Handling file upload")
    word_file = request.files.get('wordfile')
    excel_file = request.files.get('excelfile')

    if not word_file or not allowed_file(word_file.filename, ALLOWED_WORD_EXTENSIONS):
        logging.error("Invalid or missing Word (.docx) file")
        abort(400, 'Invalid or missing Word (.docx) file.')

    if not excel_file or not allowed_file(excel_file.filename, ALLOWED_EXCEL_EXTENSIONS):
        logging.error("Invalid or missing Excel (.xlsx or .xls) file")
        abort(400, 'Invalid or missing Excel (.xlsx or .xls) file.')

    word_filename = secure_filename(word_file.filename)
    excel_filename = secure_filename(excel_file.filename)

    word_filepath = os.path.join(app.config['UPLOAD_FOLDER'], word_filename)
    excel_filepath = os.path.join(app.config['UPLOAD_FOLDER'], excel_filename)

    try:
        word_file.save(word_filepath)
        excel_file.save(excel_filepath)
    except Exception as e:
        logging.error(f"Failed to save files: {str(e)}")
        abort(500, 'Failed to save uploaded files.')

    try:
        data = pd.read_excel(excel_filepath)
        logging.info(f"Excel file read successfully. Columns: {data.columns.tolist()}")
    except Exception as e:
        logging.error(f"Failed to read Excel file: {str(e)}")
        abort(400, 'Failed to read the Excel file. Ensure it is a valid format.')

    columns = data.columns.tolist()
    return render_template('choose_column.html', columns=columns, word_filepath=word_filepath, excel_filepath=excel_filepath)

@app.route('/process', methods=['POST'])
def process_files():
    logging.info("Processing files")
    word_filepath = request.form['word_filepath']
    excel_filepath = request.form['excel_filepath']
    chosen_column = request.form['chosen_column']

    try:
        data = pd.read_excel(excel_filepath, engine='openpyxl')
        logging.info(f"Excel columns: {data.columns.tolist()}")
    except Exception as e:
        logging.error(f"Error reading Excel file during processing: {str(e)}")
        abort(400, 'Error reading the Excel file during processing.')

    if chosen_column not in data.columns:
        logging.error(f"Chosen column '{chosen_column}' not found in Excel file")
        abort(400, f"Column '{chosen_column}' not found in Excel file.")

    filenames = []

    for index, row in data.iterrows():
        try:
            doc = Document(word_filepath)
        except Exception as e:
            logging.error(f"Error loading Word document: {str(e)}")
            continue

        # Create a dictionary mapping from column names to values
        row_dict = {}
        for col in data.columns:
            value = row[col]
            if isinstance(value, datetime.datetime) and pd.notna(value):
                value = value.strftime('%d/%m/%Y')
            elif pd.isna(value):
                value = ""
            else:
                value = str(value).strip()
            row_dict[col] = value

        # Log all possible placeholders with column names for debugging
        logging.info(f"Row data keys: {list(row_dict.keys())}")
        logging.info(f"Row data: {row_dict}")

        # Log all tables and placeholders
        placeholders_found = set()
        for table in doc.tables:
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    for paragraph in cell.paragraphs:
                        full_text = ''.join(run.text for run in paragraph.runs)
                        normalized_text = re.sub(r'[\s\u00a0\u200b\u00ad\u200c\u200d\u2028\u200e\u200f]+', ' ', full_text).strip()
                        run_texts = [repr(run.text) for run in paragraph.runs]
                        unicode_chars = [(char, hex(ord(char))) for char in normalized_text if ord(char) > 127 or ord(char) < 32]
                        matches = re.findall(r"[\u00ab\u2039<](.*?[\u00bb\u203a>])", normalized_text)
                        placeholders_found.update(matches)
                        logging.debug(f"Table cell [row {row_idx}, col {col_idx}] runs: {run_texts}")
                        logging.debug(f"Table cell [row {row_idx}, col {col_idx}] unicode chars: {unicode_chars}")
                        logging.debug(f"Table cell [row {row_idx}, col {col_idx}] full text: {full_text!r}, normalized: {normalized_text!r}, matches: {matches}")

        for para_idx, paragraph in enumerate(doc.paragraphs):
            full_text = ''.join(run.text for run in paragraph.runs)
            normalized_text = re.sub(r'[\s\u00a0\u200b\u00ad\u200c\u200d\u2028\u200e\u200f]+', ' ', full_text).strip()
            run_texts = [repr(run.text) for run in paragraph.runs]
            unicode_chars = [(char, hex(ord(char))) for char in normalized_text if ord(char) > 127 or ord(char) < 32]
            matches = re.findall(r"[\u00ab\u2039<](.*?[\u00bb\u203a>])", normalized_text)
            placeholders_found.update(matches)
            logging.debug(f"Paragraph {para_idx} runs: {run_texts}")
            logging.debug(f"Paragraph {para_idx} unicode chars: {unicode_chars}")
            logging.debug(f"Paragraph {para_idx} full text: {full_text!r}, normalized: {normalized_text!r}, matches: {matches}")
        logging.info(f"Placeholders found in document: {placeholders_found}")

        # Replace placeholders in paragraphs and tables
        total_replacements = 0
        for para_idx, paragraph in enumerate(doc.paragraphs):
            replacements = _replace_placeholders_in_paragraph(paragraph, row_dict, para_idx)
            total_replacements += replacements

        for table_idx, table in enumerate(doc.tables):
            replacements = _replace_placeholders_in_table(table, row_dict, table_idx)
            total_replacements += replacements

        logging.info(f"Total replacements made: {total_replacements}")
        if total_replacements == 0:
            logging.warning("No placeholders were replaced in the document")

        try:
            output_filename = f"{secure_filename(str(row_dict[chosen_column]))}_{index}.docx"
            output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)
            doc.save(output_path)
            filenames.append(output_path)
        except KeyError as e:
            logging.error(f"Column '{chosen_column}' not found in row data: {str(e)}")
            continue
        except Exception as e:
            logging.error(f"Error saving output file {output_filename}: {str(e)}")
            continue

    if not filenames:
        logging.error("No files were generated")
        abort(500, 'No files were generated due to processing errors.')

    zip_path = os.path.join(app.config['OUTPUT_FOLDER'], "processed_documents.zip")

    try:
        with zipfile.ZipFile(zip_path, 'w') as doc_zip:
            for file in filenames:
                doc_zip.write(file, arcname=os.path.basename(file))
        logging.info(f"Created zip file: {zip_path}")
    except Exception as e:
        logging.error(f"Error creating zip file: {str(e)}")
        abort(500, 'Error creating zip file.')

    @after_this_request
    def remove_files(response):
        logging.info("Cleaning up temporary files")
        try:
            for file in filenames:
                if os.path.exists(file):
                    os.remove(file)
            for file in [word_filepath, excel_filepath]:
                if os.path.exists(file):
                    os.remove(file)
            if os.path.exists(zip_path):
                os.remove(zip_path)
        except Exception as e:
            logging.error(f"Error cleaning up files: {str(e)}")
        return response

    logging.info("Sending zip file to client")
    return send_from_directory(app.config['OUTPUT_FOLDER'], "processed_documents.zip", as_attachment=True)

def _replace_placeholders_in_paragraph(paragraph, row_data, para_idx):
    # Concatenate all runs to get the full paragraph text
    full_text = ''.join(run.text for run in paragraph.runs)
    # Enhanced normalization for Unicode whitespace and non-printable characters
    normalized_text = re.sub(r'[\s\u00a0\u200b\u00ad\u200c\u200d\u2028\u200e\u200f]+', ' ', full_text).strip()
    original_text = normalized_text
    
    # Check if there's any text to process
    if not normalized_text:
        return 0
    
    replacements = 0
    logging.debug(f"Processing paragraph {para_idx} text: {full_text!r}, normalized: {normalized_text!r}")
    
    # First approach: Direct placeholder replacement (exact column names)
    for column_name, value in row_data.items():
        # Create pattern for various placeholder formats with the column name
        patterns = [
            f"[\u00ab\u2039<]\\s*{re.escape(column_name)}\\s*[\u00bb\u203a>]",           # Basic: «Column_Name»
            f"[\u00ab\u2039<]\\s*{re.escape(column_name.lower())}\\s*[\u00bb\u203a>]",   # Lowercase: «column_name»
            f"[\u00ab\u2039<]\\s*{re.escape(column_name.upper())}\\s*[\u00bb\u203a>]"    # Uppercase: «COLUMN_NAME»
        ]
        
        for pattern in patterns:
            regex = re.compile(pattern, re.IGNORECASE)
            if regex.search(normalized_text):
                normalized_text = regex.sub(value, normalized_text)
                replacements += 1
                logging.debug(f"Replaced '{pattern}' with '{value}' in paragraph {para_idx}")
    
    # Second approach: Extract placeholders and try to match them to column names
    if replacements == 0:
        placeholders = re.findall(r"[\u00ab\u2039<](.*?[\u00bb\u203a>])", original_text)
        for placeholder in placeholders:
            # Remove the closing bracket from the placeholder
            clean_placeholder = placeholder[:-1].strip()
            # Try different variations of the placeholder to match column names
            for column_name, value in row_data.items():
                if (clean_placeholder.lower() == column_name.lower() or
                    clean_placeholder.lower().replace("_", "") == column_name.lower().replace("_", "") or
                    clean_placeholder.lower().replace(" ", "") == column_name.lower().replace(" ", "")):
                    
                    pattern = f"[\u00ab\u2039<]\\s*{re.escape(clean_placeholder)}\\s*[\u00bb\u203a>]"
                    regex = re.compile(pattern, re.IGNORECASE)
                    if regex.search(normalized_text):
                        normalized_text = regex.sub(value, normalized_text)
                        replacements += 1
                        logging.debug(f"Fuzzy match: Replaced '{pattern}' with '{value}' in paragraph {para_idx}")
    
    # If we made replacements using regex approaches, update the paragraph
    if replacements > 0:
        # Clear all runs and add a single new run with the replaced text
        for run in paragraph.runs:
            run.text = ''
        paragraph.add_run(normalized_text)
        logging.debug(f"Updated paragraph {para_idx} text: {normalized_text!r}")
    
    # Third approach: Direct run-by-run inspection and replacement
    # This is needed because some placeholders might be split across runs
    if replacements == 0:
        # Check each run for placeholders
        for i, run in enumerate(paragraph.runs):
            for column_name, value in row_data.items():
                # Different variations of placeholder formats to check
                placeholders = [
                    f"«{column_name}»", 
                    f"«{column_name.lower()}»",
                    f"«{column_name.upper()}»",
                    f"\u00ab{column_name}\u00bb",
                    f"\u00ab{column_name.lower()}\u00bb",
                    f"\u00ab{column_name.upper()}\u00bb"
                ]
                
                for placeholder in placeholders:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, value)
                        replacements += 1
                        logging.debug(f"Direct run replacement: '{placeholder}' with '{value}' in run {i} of paragraph {para_idx}")
    
    return replacements

def _replace_placeholders_in_table(table, row_data, table_idx):
    replacements = 0
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                cell_replacements = _replace_placeholders_in_paragraph(paragraph, row_data, f"table_{table_idx}_cell_{row_idx}_{col_idx}")
                replacements += cell_replacements
            # Handle nested tables if any
            for i, nested_table in enumerate(cell._element.xpath('.//w:tbl')):
                if i > 0:  # Skip the first one as it's the table itself
                    try:
                        nested_table_obj = table.__class__(nested_table, table._parent)
                        replacements += _replace_placeholders_in_table(nested_table_obj, row_data, f"{table_idx}_nested_{i}")
                    except Exception as e:
                        logging.error(f"Error processing nested table: {str(e)}")
    return replacements

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 10000))
    app.run(host='0.0.0.0', port=port, debug=False)
